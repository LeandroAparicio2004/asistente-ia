import os
import json
import csv
from config import SANDBOX_PATH

# Intentar importar docx (opcional)
try:
    from docx import Document as DocxDocument
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False

# ─── CONTEXTO ACTIVO ──────────────────────────────────────────────────────────

_archivo_activo = None

def establecer_archivo_activo(nombre: str) -> str:
    global _archivo_activo
    nombre_real = _resolver_nombre(nombre)
    ruta = os.path.abspath(os.path.join(SANDBOX_PATH, nombre_real))
    if not ruta.startswith(os.path.abspath(SANDBOX_PATH)):
        return "❌ Acceso fuera del sandbox bloqueado."
    if not os.path.exists(ruta):
        return f"❌ El archivo '{nombre}' no existe en el sandbox."
    _archivo_activo = nombre_real
    return f"📂 Ahora estamos trabajando con '{nombre_real}'. ¿Qué querés hacer con él?"

def salir_archivo_activo() -> str:
    global _archivo_activo
    if not _archivo_activo:
        return "ℹ️ No había ningún archivo activo."
    nombre = _archivo_activo
    _archivo_activo = None
    return f"✅ Saliste del archivo '{nombre}'."

def obtener_archivo_activo() -> str | None:
    return _archivo_activo

# ─── UTILIDADES ───────────────────────────────────────────────────────────────

def _resolver_nombre(nombre: str) -> str:
    ruta = os.path.abspath(os.path.join(SANDBOX_PATH, nombre))
    if os.path.exists(ruta):
        return nombre
    directorio = os.path.dirname(ruta)
    nombre_base = os.path.basename(nombre)
    if os.path.exists(directorio):
        for item in os.listdir(directorio):
            if os.path.splitext(item)[0].lower() == nombre_base.lower():
                return os.path.join(os.path.dirname(nombre), item) if os.path.dirname(nombre) else item
    return nombre

def _ruta_segura(nombre: str) -> str:
    ruta = os.path.abspath(os.path.join(SANDBOX_PATH, nombre))
    if not ruta.startswith(os.path.abspath(SANDBOX_PATH)):
        raise ValueError("❌ Acceso fuera del sandbox bloqueado.")
    return ruta

def _extension(nombre: str) -> str:
    return os.path.splitext(nombre)[1].lower()

# ─── LEER ─────────────────────────────────────────────────────────────────────

def leer_archivo_completo(nombre: str = "") -> str:
    objetivo = nombre or _archivo_activo
    if not objetivo:
        return "❌ No hay archivo activo. Decí 'Trabajemos con [nombre]' primero."
    objetivo = _resolver_nombre(objetivo)
    ruta = _ruta_segura(objetivo)
    if not os.path.exists(ruta):
        return f"❌ '{objetivo}' no existe."
    ext = _extension(objetivo)
    try:
        if ext == ".docx":
            return _leer_docx(ruta, objetivo)
        elif ext == ".json":
            return _leer_json(ruta, objetivo)
        elif ext == ".csv":
            return _leer_csv(ruta, objetivo)
        else:
            with open(ruta, "r", encoding="utf-8") as f:
                contenido = f.read()
            if not contenido.strip():
                return f"📄 '{objetivo}' está vacío."
            return f"📄 Contenido de '{objetivo}':\n\n{contenido}"
    except Exception as e:
        return f"❌ Error al leer '{objetivo}': {str(e)}"

def _leer_docx(ruta: str, nombre: str) -> str:
    if not DOCX_DISPONIBLE:
        return "❌ Necesitás instalar python-docx: pip install python-docx"
    doc = DocxDocument(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not texto:
        return f"📄 '{nombre}' está vacío."
    return f"📄 Contenido de '{nombre}':\n\n{texto}"

def _leer_json(ruta: str, nombre: str) -> str:
    with open(ruta, "r", encoding="utf-8") as f:
        data = json.load(f)
    return f"📄 Contenido de '{nombre}':\n\n{json.dumps(data, ensure_ascii=False, indent=2)}"

def _leer_csv(ruta: str, nombre: str) -> str:
    with open(ruta, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        filas = list(reader)
    if not filas:
        return f"📄 '{nombre}' está vacío."
    resultado = f"📄 Contenido de '{nombre}' ({len(filas)} filas):\n\n"
    for fila in filas[:50]:
        resultado += "  " + " | ".join(fila) + "\n"
    if len(filas) > 50:
        resultado += f"\n  ... y {len(filas) - 50} filas más."
    return resultado

# ─── ESCRIBIR / EDITAR ────────────────────────────────────────────────────────

def agregar_contenido(texto: str = "", nombre: str = "", **kwargs) -> str:
    # Aceptar parámetros alternativos
    texto = texto or kwargs.get("contenido") or kwargs.get("archivo") or ""
    objetivo = nombre or _archivo_activo
    if not objetivo:
        return "❌ No hay archivo activo. Decí 'Trabajemos con [nombre]' primero."
    if not texto:
        return "❌ No especificaste qué escribir."
    objetivo = _resolver_nombre(objetivo)
    ruta = _ruta_segura(objetivo)
    if not os.path.exists(ruta):
        return f"❌ '{objetivo}' no existe."
    ext = _extension(objetivo)
    try:
        if ext == ".docx":
            return _agregar_docx(ruta, objetivo, texto)
        elif ext == ".json":
            return "⚠️ Para modificar JSON usá 'reemplazar contenido'."
        else:
            with open(ruta, "a", encoding="utf-8") as f:
                f.write("\n" + texto)
            return f"✅ Texto agregado al final de '{objetivo}'."
    except Exception as e:
        return f"❌ Error: {str(e)}"

def _agregar_docx(ruta: str, nombre: str, texto: str) -> str:
    if not DOCX_DISPONIBLE:
        return "❌ Necesitás instalar python-docx: pip install python-docx"
    doc = DocxDocument(ruta)
    doc.add_paragraph(texto)
    doc.save(ruta)
    return f"✅ Texto agregado al final de '{nombre}'."

def reemplazar_contenido(contenido_nuevo: str, nombre: str = "") -> str:
    objetivo = nombre or _archivo_activo
    if not objetivo:
        return "❌ No hay archivo activo."
    objetivo = _resolver_nombre(objetivo)
    ruta = _ruta_segura(objetivo)
    if not os.path.exists(ruta):
        return f"❌ '{objetivo}' no existe."
    ext = _extension(objetivo)
    try:
        if ext == ".docx":
            if not DOCX_DISPONIBLE:
                return "❌ Necesitás instalar python-docx"
            doc = DocxDocument(ruta)
            for p in doc.paragraphs:
                p.clear()
            doc.paragraphs[0].text = contenido_nuevo
            doc.save(ruta)
        else:
            with open(ruta, "w", encoding="utf-8") as f:
                f.write(contenido_nuevo)
        return f"✅ Contenido de '{objetivo}' reemplazado."
    except Exception as e:
        return f"❌ Error: {str(e)}"

def borrar_linea(numero_linea: int, nombre: str = "") -> str:
    objetivo = nombre or _archivo_activo
    if not objetivo:
        return "❌ No hay archivo activo."
    objetivo = _resolver_nombre(objetivo)
    ruta = _ruta_segura(objetivo)
    if not os.path.exists(ruta):
        return f"❌ '{objetivo}' no existe."
    ext = _extension(objetivo)
    if ext == ".docx":
        return "⚠️ Para Word usá 'reemplazar contenido'."
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            lineas = f.readlines()
        if numero_linea < 1 or numero_linea > len(lineas):
            return f"❌ La línea {numero_linea} no existe. El archivo tiene {len(lineas)} líneas."
        linea_borrada = lineas.pop(numero_linea - 1).strip()
        with open(ruta, "w", encoding="utf-8") as f:
            f.writelines(lineas)
        return f"✅ Línea {numero_linea} borrada: '{linea_borrada}'"
    except Exception as e:
        return f"❌ Error: {str(e)}"

def buscar_en_archivo(termino: str, nombre: str = "") -> str:
    objetivo = nombre or _archivo_activo
    if not objetivo:
        return "❌ No hay archivo activo."
    objetivo = _resolver_nombre(objetivo)
    ruta = _ruta_segura(objetivo)
    if not os.path.exists(ruta):
        return f"❌ '{objetivo}' no existe."
    ext = _extension(objetivo)
    try:
        if ext == ".docx":
            if not DOCX_DISPONIBLE:
                return "❌ Necesitás instalar python-docx"
            doc = DocxDocument(ruta)
            lineas = [p.text for p in doc.paragraphs]
        else:
            with open(ruta, "r", encoding="utf-8") as f:
                lineas = f.readlines()
        resultados = []
        for i, linea in enumerate(lineas, 1):
            if termino.lower() in linea.lower():
                resultados.append(f"  Línea {i}: {linea.strip()}")
        if not resultados:
            return f"🔍 No encontré '{termino}' en '{objetivo}'."
        return f"🔍 '{termino}' encontrado {len(resultados)} vez/veces:\n" + "\n".join(resultados)
    except Exception as e:
        return f"❌ Error: {str(e)}"

def generar_y_escribir(tema: str, nombre_archivo: str, formato: str = "txt") -> str:
    try:
        from groq import Groq
        from config import GROQ_API_KEY
        import os

        # Asegurar extensión correcta
        if not any(nombre_archivo.endswith(ext) for ext in [".txt", ".md", ".docx"]):
            nombre_archivo = f"{nombre_archivo}.{formato}"

        ruta = _ruta_segura(nombre_archivo)

        print(f"   ✍️  Generando contenido sobre '{tema}'...")

        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "Sos un asistente experto en redacción. Cuando te pidan escribir sobre un tema, generá un documento completo, bien estructurado, con títulos, subtítulos y contenido detallado. Escribí en español. No uses markdown con asteriscos, usá texto plano con MAYÚSCULAS para los títulos."
                },
                {
                    "role": "user",
                    "content": f"Escribí un informe completo y detallado sobre: {tema}"
                }
            ],
            temperature=0.7,
            max_tokens=2000
        )

        contenido = response.choices[0].message.content.strip()
        ext = os.path.splitext(nombre_archivo)[1].lower()

        if ext == ".docx":
            if not DOCX_DISPONIBLE:
                return "❌ Necesitás instalar python-docx: pip install python-docx"
            from docx import Document as DocxDocument
            doc = DocxDocument()
            doc.add_heading(tema, 0)
            for linea in contenido.split("\n"):
                if linea.strip():
                    doc.add_paragraph(linea)
            doc.save(ruta)
        else:
            with open(ruta, "w", encoding="utf-8") as f:
                f.write(contenido)

        tamaño = len(contenido.split())
        return f"✅ Informe sobre '{tema}' generado en '{nombre_archivo}' ({tamaño} palabras aprox.)"

    except Exception as e:
        return f"❌ Error al generar informe: {str(e)}"
